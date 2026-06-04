#!/usr/bin/env python3
"""Minimal regression suite for paper-figure-polisher skill scripts."""

import base64
import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent.parent
SCRIPTS = ROOT / "scripts"
FIXTURES = Path(__file__).resolve().parent / "fixtures"
PYTHON = sys.executable


def run_command(args, cwd=None):
    result = subprocess.run(
        args,
        cwd=str(cwd or ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    if result.returncode != 0:
        raise AssertionError("Command failed:\n%s" % result.stdout)
    return result.stdout


def assert_equal(actual, expected, message):
    if actual != expected:
        raise AssertionError("%s: expected %r, got %r" % (message, expected, actual))


def assert_true(value, message):
    if not value:
        raise AssertionError(message)


def make_png(path):
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9WnR6WQAAAAASUVORK5CYII="
    )
    path.write_bytes(png)


def test_baseline_table_conversion(tmpdir):
    src = FIXTURES / "baseline_simple.csv"
    out = tmpdir / "baseline_simple.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "spreadsheet_to_three_line_table.py"),
            str(src),
            str(out),
            "--caption",
            "Table 1.Baseline sensitivity analysis",
        ]
    )
    doc = Document(str(out))
    table = doc.tables[0]
    assert_equal(len(table.rows), 4, "baseline table row count")
    assert_equal(table.rows[1].cells[0].text, "Unadjusted", "baseline first data row")


def test_insert_new_table_before_existing_figures(tmpdir):
    src = FIXTURES / "baseline_simple.csv"
    img = tmpdir / "dot.png"
    make_png(img)

    manuscript = tmpdir / "manuscript_with_table_then_figure.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph("Table 1.Existing baseline table")
    existing_table = doc.add_table(rows=2, cols=2)
    existing_table.rows[0].cells[0].text = "Item"
    existing_table.rows[0].cells[1].text = "Value"
    existing_table.rows[1].cells[0].text = "Age"
    existing_table.rows[1].cells[1].text = "65"
    doc.add_paragraph("Note. Existing baseline note.")
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("Figure 1.Existing workflow figure")
    doc.save(str(manuscript))

    out = tmpdir / "manuscript_with_inserted_table.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "spreadsheet_to_three_line_table.py"),
            str(src),
            str(out),
            "--input-docx",
            str(manuscript),
            "--caption",
            "Table 2.Inserted sensitivity analysis",
        ]
    )

    out_doc = Document(str(out))
    non_empty_paragraphs = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert_equal(non_empty_paragraphs[0], "Intro paragraph.", "intro paragraph should remain first")
    assert_equal(non_empty_paragraphs[1], "Table 1.Existing baseline table", "existing table caption should remain first")
    assert_equal(non_empty_paragraphs[2], "Note. Existing baseline note.", "existing table note should stay with first table")
    assert_equal(non_empty_paragraphs[3], "Table 2.Inserted sensitivity analysis", "new table should be inserted after existing table block")
    assert_equal(non_empty_paragraphs[4], "Figure 1.Existing workflow figure", "figure caption should remain after all table blocks")
    assert_equal(len(out_doc.tables), 2, "document should contain existing and inserted tables")


def test_insert_new_figure_after_existing_figures(tmpdir):
    img = tmpdir / "dot.png"
    make_png(img)

    manuscript = tmpdir / "manuscript_with_table_and_figure.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph("Table 1.Existing baseline table")
    existing_table = doc.add_table(rows=2, cols=2)
    existing_table.rows[0].cells[0].text = "Item"
    existing_table.rows[0].cells[1].text = "Value"
    existing_table.rows[1].cells[0].text = "Age"
    existing_table.rows[1].cells[1].text = "65"
    doc.add_paragraph("Note. Existing baseline note.")
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("Figure 1.Existing workflow figure")
    doc.add_paragraph("Note. Existing figure note.")
    doc.save(str(manuscript))

    out = tmpdir / "manuscript_with_inserted_figure.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "insert_figures_into_docx.py"),
            str(manuscript),
            str(out),
            str(img),
            "--captions",
            "New outcome flowchart",
        ]
    )

    out_doc = Document(str(out))
    non_empty_paragraphs = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert_equal(non_empty_paragraphs[1], "Table 1.Existing baseline table", "table caption should remain before figures")
    assert_equal(non_empty_paragraphs[3], "Figure 1.Existing workflow figure", "existing figure caption should remain first")
    assert_equal(non_empty_paragraphs[4], "Note. Existing figure note.", "existing figure note should remain attached")
    assert_equal(non_empty_paragraphs[5], "Figure 2.New outcome flowchart", "new figure should be inserted after existing figure block")


def test_insert_table_before_figures_when_no_existing_tables(tmpdir):
    src = FIXTURES / "baseline_simple.csv"
    img = tmpdir / "dot.png"
    make_png(img)

    manuscript = tmpdir / "manuscript_with_only_figure.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("Figure 1.Existing workflow figure")
    doc.save(str(manuscript))

    out = tmpdir / "manuscript_with_first_table_before_figure.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "spreadsheet_to_three_line_table.py"),
            str(src),
            str(out),
            "--input-docx",
            str(manuscript),
            "--caption",
            "Table 1.Inserted baseline table",
        ]
    )

    out_doc = Document(str(out))
    non_empty_paragraphs = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert_equal(non_empty_paragraphs[0], "Intro paragraph.", "intro paragraph should remain first")
    assert_equal(non_empty_paragraphs[1], "Table 1.Inserted baseline table", "first table should be inserted before the first figure")
    assert_equal(non_empty_paragraphs[2], "Figure 1.Existing workflow figure", "figure should remain after inserted table")


def test_grouped_keyword_inference(tmpdir):
    src = FIXTURES / "grouped_keywords.csv"
    out = tmpdir / "grouped_keywords.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "spreadsheet_to_three_line_table.py"),
            str(src),
            str(out),
            "--caption",
            "Table 2.Grouped keyword inference",
            "--landscape",
        ]
    )
    doc = Document(str(out))
    rows = [row.cells[0].text for row in doc.tables[0].rows]
    expected_groups = [
        "人口学与住院信息",
        "既往史与合并症",
        "基线实验室与辅助检查",
        "随访指标",
        "结局事件",
    ]
    for group in expected_groups:
        assert_true(group in rows, "missing inferred group heading %s" % group)

    rules_path = ROOT / "assets" / "medical-grouping-rules.json"
    assert_true(rules_path.exists(), "missing fixed medical grouping rules file")

    custom_rules = tmpdir / "custom_group_rules.json"
    custom_rules.write_text(
        json.dumps(
            {
                "default_group": "其他项目",
                "followup_group": "统一随访",
                "outcome_group": "统一结局",
                "preserve_previous_suffixes": [":", "："],
                "followup_markers": ["术后1个月", "术后3个月", "术后6/9个月", "术后12个月"],
                "outcome_keywords": ["死亡", "心肌梗死", "卒中", "出血", "事件"],
                "outcome_forced_labels": ["死亡时间"],
                "groups": [
                    {"title": "自定义人口学", "keywords": ["住院号", "性别", "年龄"]},
                    {"title": "自定义合并症", "keywords": ["高血压", "糖尿病"]},
                    {"title": "自定义基线", "keywords": ["基线", "wbc", "ast"]}
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    custom_out = tmpdir / "grouped_custom.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "spreadsheet_to_three_line_table.py"),
            str(src),
            str(custom_out),
            "--caption",
            "Table 3.Custom grouping rules",
            "--group-rules",
            str(custom_rules),
        ]
    )
    custom_doc = Document(str(custom_out))
    custom_rows = [row.cells[0].text for row in custom_doc.tables[0].rows]
    assert_true("自定义人口学" in custom_rows, "custom grouping rules were not applied")
    assert_true("统一随访" in custom_rows, "custom follow-up group was not applied")
    assert_true("统一结局" in custom_rows, "custom outcome group was not applied")


def test_docx_table_caption_cleanup_and_autogen(tmpdir):
    dirty_doc = tmpdir / "table_dirty.docx"
    doc = Document()
    doc.add_paragraph("table1 : summary of experimental conditions")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Age"
    table.rows[1].cells[1].text = "65"
    doc.save(str(dirty_doc))

    cleaned = tmpdir / "table_dirty_out.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "normalize_docx_tables.py"),
            str(dirty_doc),
            str(cleaned),
        ]
    )
    cleaned_doc = Document(str(cleaned))
    assert_equal(
        cleaned_doc.paragraphs[0].text,
        "Table 1.summary of experimental conditions",
        "dirty table caption cleanup",
    )

    missing_doc = tmpdir / "table_missing.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Age"
    table.rows[1].cells[1].text = "65"
    doc.add_paragraph("Note. Existing table note.")
    doc.save(str(missing_doc))

    generated = tmpdir / "table_missing_out.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "normalize_docx_tables.py"),
            str(missing_doc),
            str(generated),
        ]
    )
    generated_doc = Document(str(generated))
    assert_equal(generated_doc.paragraphs[0].text, "Table 1.", "missing table caption autogen")
    assert_equal(
        generated_doc.paragraphs[1].text,
        "Note. Existing table note.",
        "table note preserved after caption insertion",
    )


def test_complex_docx_table_normalization(tmpdir):
    src = tmpdir / "complex_tables.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph before the first table.")
    doc.add_paragraph("table1 : baseline characteristics")
    table1 = doc.add_table(rows=3, cols=3)
    table1.rows[0].cells[0].text = "Item"
    table1.rows[0].cells[1].text = "Value A"
    table1.rows[0].cells[2].text = "Value B"
    merged = table1.rows[1].cells[0].merge(table1.rows[1].cells[1])
    merged = merged.merge(table1.rows[1].cells[2])
    merged.text = "人口学与住院信息"
    table1.rows[2].cells[0].text = "年龄"
    table1.rows[2].cells[1].text = "65"
    table1.rows[2].cells[2].text = "64"
    doc.add_paragraph("Note. Existing note for table 1.")
    doc.add_paragraph("Body paragraph between tables should remain unchanged.")
    table2 = doc.add_table(rows=2, cols=6)
    headers = ["Var", "All", "A", "B", "OR", "P"]
    for i, header in enumerate(headers):
        table2.rows[0].cells[i].text = header
        table2.rows[1].cells[i].text = "Very long content %s" % i
    doc.add_paragraph("Note. Existing note for table 2.")
    table3 = doc.add_table(rows=2, cols=2)
    table3.rows[0].cells[0].text = "Item"
    table3.rows[0].cells[1].text = "Value"
    table3.rows[1].cells[0].text = "BMI"
    table3.rows[1].cells[1].text = "24.1"
    doc.add_paragraph("Note. Existing note for table 3.")
    doc.save(str(src))

    out = tmpdir / "complex_tables_out.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "normalize_docx_tables.py"),
            str(src),
            str(out),
        ]
    )
    out_doc = Document(str(out))
    non_empty_paragraphs = [p.text.strip() for p in out_doc.paragraphs if p.text.strip()]
    assert_equal(non_empty_paragraphs[0], "Intro paragraph before the first table.", "intro paragraph preserved")
    assert_equal(non_empty_paragraphs[1], "Table 1.baseline characteristics", "first table caption cleaned")
    assert_equal(non_empty_paragraphs[2], "Note. Existing note for table 1.", "first table note preserved")
    assert_equal(
        non_empty_paragraphs[3],
        "Body paragraph between tables should remain unchanged.",
        "intermediate body paragraph preserved",
    )
    assert_equal(non_empty_paragraphs[4], "Table 2.", "second table caption auto-generated")
    assert_equal(non_empty_paragraphs[5], "Note. Existing note for table 2.", "second table note preserved")
    assert_equal(non_empty_paragraphs[6], "Table 3.", "third narrow table caption auto-generated")
    assert_equal(non_empty_paragraphs[7], "Note. Existing note for table 3.", "third narrow table note preserved")
    orientations = [section.orientation for section in out_doc.sections]
    assert_equal(
        orientations,
        [WD_ORIENT.PORTRAIT, WD_ORIENT.LANDSCAPE, WD_ORIENT.PORTRAIT],
        "only the wide table block should be isolated into a landscape section",
    )
    merged_heading = out_doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0].bold
    assert_true(merged_heading, "merged heading row should stay bold")
    assert_true(out_doc.tables[0].rows[0]._tr.xml.find("tblHeader") != -1, "header row should repeat across pages")


def test_docx_figure_caption_cleanup_and_autogen(tmpdir):
    img = tmpdir / "dot.png"
    make_png(img)

    dirty_doc = tmpdir / "figure_dirty.docx"
    doc = Document()
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("figure1 schematic diagram of the 3D canopy reconstruction process")
    doc.save(str(dirty_doc))

    cleaned = tmpdir / "figure_dirty_out.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "paginate_docx_figures.py"),
            str(dirty_doc),
            str(cleaned),
        ]
    )
    cleaned_doc = Document(str(cleaned))
    figure_paragraphs = [p.text.strip() for p in cleaned_doc.paragraphs if p.text.strip()]
    assert_equal(
        figure_paragraphs[0],
        "Figure 1.schematic diagram of the 3D canopy reconstruction process",
        "dirty figure caption cleanup",
    )

    missing_doc = tmpdir / "figure_missing.docx"
    doc = Document()
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("Note. Existing figure note.")
    doc.add_picture(str(img), width=Inches(1.0))
    doc.add_paragraph("Some trailing body text.")
    doc.save(str(missing_doc))

    generated = tmpdir / "figure_missing_out.docx"
    run_command(
        [
            PYTHON,
            str(SCRIPTS / "paginate_docx_figures.py"),
            str(missing_doc),
            str(generated),
        ]
    )
    generated_doc = Document(str(generated))
    figure_paragraphs = [p.text.strip() for p in generated_doc.paragraphs if p.text.strip()]
    assert_equal(figure_paragraphs[0], "Figure 1.", "first missing figure caption autogen")
    assert_equal(figure_paragraphs[1], "Note. Existing figure note.", "figure note preserved")
    assert_equal(figure_paragraphs[2], "Figure 2.", "second missing figure caption autogen")


def main():
    temp_root = Path(tempfile.mkdtemp(prefix="paper-figure-polisher-tests-"))
    try:
        test_baseline_table_conversion(temp_root)
        test_insert_new_table_before_existing_figures(temp_root)
        test_insert_new_figure_after_existing_figures(temp_root)
        test_insert_table_before_figures_when_no_existing_tables(temp_root)
        test_grouped_keyword_inference(temp_root)
        test_docx_table_caption_cleanup_and_autogen(temp_root)
        test_complex_docx_table_normalization(temp_root)
        test_docx_figure_caption_cleanup_and_autogen(temp_root)
        print("All regression tests passed.")
    finally:
        shutil.rmtree(str(temp_root), ignore_errors=True)


if __name__ == "__main__":
    main()
