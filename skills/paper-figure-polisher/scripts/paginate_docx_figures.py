#!/usr/bin/env python3
"""Insert page breaks and normalize figure captions for DOCX figures."""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


FIGURE_CAPTION_RE = re.compile(r"^(figure|图)\s*(\d+)\s*[\.:\-：]?\s*", re.IGNORECASE)
NOTE_RE = re.compile(r"^(note\.|注[:：])", re.IGNORECASE)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Insert page breaks before DOCX figure paragraphs."
    )
    parser.add_argument("source", help="Path to source .docx")
    parser.add_argument("output", help="Path to output .docx")
    parser.add_argument(
        "--break-before-first-figure",
        action="store_true",
        help="Also insert a page break before the first detected figure",
    )
    parser.add_argument(
        "--font-name",
        default="Times New Roman",
        help="Western font for captions; defaults to Times New Roman",
    )
    parser.add_argument(
        "--east-asia-font",
        default="宋体",
        help="East Asian body font; defaults to 宋体",
    )
    parser.add_argument(
        "--east-asia-heading-font",
        default="黑体",
        help="East Asian heading font for Chinese caption labels; defaults to 黑体",
    )
    parser.add_argument(
        "--caption-size",
        default=11.0,
        type=float,
        help="Figure caption font size in points; defaults to 11.0",
    )
    parser.add_argument(
        "--note-size",
        default=10.0,
        type=float,
        help="Figure note font size in points; defaults to 10.0",
    )
    return parser.parse_args()


def has_drawing(paragraph):
    for node in paragraph._p.iter():
        if node.tag.endswith("}drawing") or node.tag.endswith("}pict"):
            return True
    return False


def paragraph_is_page_break(paragraph):
    for node in paragraph._p.iter():
        if node.tag.endswith("}br") and node.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
        ) == "page":
            return True
    return False


def previous_paragraph(paragraph):
    sibling = paragraph._p.getprevious()
    while sibling is not None:
        if sibling.tag.endswith("}p"):
            for item in paragraph._parent.paragraphs:
                if item._p is sibling:
                    return item
        sibling = sibling.getprevious()
    return None


def needs_page_break(paragraph):
    previous = previous_paragraph(paragraph)
    if previous is None:
        return False
    return not paragraph_is_page_break(previous)


def insert_page_break_before(paragraph):
    before = paragraph.insert_paragraph_before()
    before.add_run().add_break(WD_BREAK.PAGE)


def is_chinese_text(text):
    return any("\u4e00" <= ch <= "\u9fff" for ch in str(text))


def set_run_fonts(run, western_font, east_asia_font, size, bold=False, italic=False):
    run.font.name = western_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western_font)
    r_fonts.set(qn("w:hAnsi"), western_font)
    r_fonts.set(qn("w:cs"), western_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def style_caption_paragraph(
    paragraph,
    western_font,
    east_asia_font,
    size,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        east_font = east_asia_font if is_chinese_text(run.text) else east_asia_font
        set_run_fonts(
            run,
            western_font,
            east_font,
            size,
            bold=bold,
            italic=italic,
        )


def caption_tail(text, pattern):
    stripped = str(text).strip()
    match = pattern.match(stripped)
    if not match:
        return ""
    return stripped[match.end():].strip()


def extract_caption_number(text, pattern):
    stripped = str(text).strip()
    match = pattern.match(stripped)
    if not match:
        return None
    try:
        return int(match.group(2))
    except Exception:
        return None


def normalize_caption_text(label, number, existing_text, pattern):
    tail = caption_tail(existing_text, pattern)
    if tail:
        return "%s %s.%s" % (label, number, tail)
    return "%s %s." % (label, number)


def next_paragraph(paragraph):
    sibling = paragraph._p.getnext()
    while sibling is not None:
        if sibling.tag.endswith("}p"):
            for item in paragraph._parent.paragraphs:
                if item._p is sibling:
                    return item
        sibling = sibling.getnext()
    return None


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def resolve_figure_caption_paragraph(figure_paragraph, figure_number):
    next_para = next_paragraph(figure_paragraph)
    if next_para and FIGURE_CAPTION_RE.match(next_para.text.strip()):
        next_para.text = normalize_caption_text("Figure", figure_number, next_para.text, FIGURE_CAPTION_RE)
        return next_para
    if next_para and NOTE_RE.match(next_para.text.strip()):
        return next_para.insert_paragraph_before("Figure %s." % figure_number)
    if next_para is not None:
        return next_para.insert_paragraph_before("Figure %s." % figure_number)
    return insert_paragraph_after(figure_paragraph, "Figure %s." % figure_number)


def style_caption_and_note(caption_paragraph, args):
    style_caption_paragraph(
        caption_paragraph,
        args.font_name,
        args.east_asia_heading_font,
        args.caption_size,
        bold=False,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    maybe_note = next_paragraph(caption_paragraph)
    if maybe_note and NOTE_RE.match(maybe_note.text.strip()):
        style_caption_paragraph(
            maybe_note,
            args.font_name,
            args.east_asia_font,
            args.note_size,
            bold=False,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )


def main():
    args = parse_args()
    document = Document(args.source)
    figure_paragraphs = [paragraph for paragraph in document.paragraphs if has_drawing(paragraph)]

    inserted = 0
    figure_number = 0
    for index, paragraph in enumerate(figure_paragraphs):
        figure_number += 1
        if index == 0 and not args.break_before_first_figure:
            if False:
                pass
        elif needs_page_break(paragraph):
            insert_page_break_before(paragraph)
            inserted += 1
        caption_paragraph = resolve_figure_caption_paragraph(paragraph, figure_number)
        style_caption_and_note(caption_paragraph, args)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print("Inserted %s page breaks into %s" % (inserted, output_path))


if __name__ == "__main__":
    main()
