#!/usr/bin/env python3
"""Normalize existing DOCX tables to the current three-line journal style."""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


TABLE_CAPTION_RE = re.compile(r"^(table|表)\s*(\d+)\s*[\.:\-：]?\s*", re.IGNORECASE)
NOTE_RE = re.compile(r"^(note\.|注[:：])", re.IGNORECASE)
NUMERIC_RE = re.compile(r"^[\s\.\-+]?[\d,]+(?:\.\d+)?(?:\s*\[.*\])?(?:\s*\(.+\))?$")
DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
PORTRAIT_CONTENT_WIDTH_CM = 16.0
A4_PORTRAIT = (11906, 16838)
A4_LANDSCAPE = (16838, 11906)
MARGIN_TWIPS = 1417


def parse_args():
    parser = argparse.ArgumentParser(
        description="Normalize existing DOCX tables to a three-line manuscript style."
    )
    parser.add_argument("source", help="Path to source .docx")
    parser.add_argument("output", help="Path to output .docx")
    parser.add_argument(
        "--font-name",
        default="Times New Roman",
        help="Western font for table text; defaults to Times New Roman",
    )
    parser.add_argument(
        "--east-asia-font",
        default="宋体",
        help="East Asian body font; defaults to 宋体",
    )
    parser.add_argument(
        "--east-asia-heading-font",
        default="黑体",
        help="East Asian heading font; defaults to 黑体",
    )
    parser.add_argument(
        "--font-size",
        default=11.0,
        type=float,
        help="Body font size in points; defaults to 11.0",
    )
    parser.add_argument(
        "--caption-size",
        default=11.0,
        type=float,
        help="Caption font size in points; defaults to 11.0",
    )
    parser.add_argument(
        "--note-size",
        default=10.0,
        type=float,
        help="Note font size in points; defaults to 10.0",
    )
    parser.add_argument(
        "--landscape",
        action="store_true",
        help="Force all sections to landscape orientation",
    )
    parser.add_argument(
        "--wide-table-section-mode",
        choices=("page", "continuous"),
        default="page",
        help=(
            "How to isolate automatically detected wide-table blocks: "
            "'page' gives the block its own page flow, 'continuous' keeps it inline"
        ),
    )
    parser.add_argument(
        "--wide-table-spacing-before-pt",
        default=6.0,
        type=float,
        help="Spacing before an isolated wide-table block in points; defaults to 6",
    )
    parser.add_argument(
        "--wide-table-spacing-after-pt",
        default=6.0,
        type=float,
        help="Spacing after an isolated wide-table block in points; defaults to 6",
    )
    return parser.parse_args()


def is_chinese_text(text):
    return any("\u4e00" <= ch <= "\u9fff" for ch in str(text))


def set_run_fonts(run, western_font, east_asia_font, size, bold=False, italic=False):
    run.font.name = western_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western_font)
    r_fonts.set(qn("w:hAnsi"), western_font)
    r_fonts.set(qn("w:cs"), western_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def style_paragraph(
    paragraph,
    western_font,
    east_asia_font,
    size,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=0,
    line_spacing=1.5,
    left_indent_cm=0.0,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        east_font = east_asia_font if is_chinese_text(run.text) else east_asia_font
        set_run_fonts(
            run,
            western_font,
            east_font,
            size,
            bold=bold,
            italic=italic,
        )


def caption_tail(text, pattern):
    stripped = str(text).strip()
    match = pattern.match(stripped)
    if not match:
        return ""
    tail = stripped[match.end():].strip()
    return tail


def extract_caption_number(text, pattern):
    stripped = str(text).strip()
    match = pattern.match(stripped)
    if not match:
        return None
    try:
        return int(match.group(2))
    except Exception:
        return None


def normalize_caption_text(label, number, existing_text, pattern):
    tail = caption_tail(existing_text, pattern)
    if tail:
        return "%s %s.%s" % (label, number, tail)
    return "%s %s." % (label, number)


def insert_paragraph_before_table(table, text=""):
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_paragraph_before_block(block, text=""):
    new_p = OxmlElement("w:p")
    element = block._p if hasattr(block, "_p") else block._tbl
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_paragraph_after_block(block, text=""):
    new_p = OxmlElement("w:p")
    element = block._p if hasattr(block, "_p") else block._tbl
    element.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if text:
        paragraph.add_run(text)
    return paragraph


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:%s" % key), str(value))


def unique_cells(cells):
    seen = set()
    result = []
    for cell in cells:
        identity = id(cell._tc)
        if identity in seen:
            continue
        seen.add(identity)
        result.append(cell)
    return result


def clear_borders(table):
    for row in table.rows:
        for cell in unique_cells(row.cells):
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )


def is_group_heading_row(row):
    texts = [cell.text.strip() for cell in row.cells]
    non_empty = [text for text in texts if text]
    if len(non_empty) == 1:
        return True
    if non_empty and len(set(non_empty)) == 1:
        return True
    return False


def is_inline_heading_row(row):
    first_text = row.cells[0].text.strip()
    if not first_text:
        return False
    if not first_text.endswith((":", "：")):
        return False
    remaining = [cell.text.strip() for cell in row.cells[1:]]
    non_empty_remaining = [text for text in remaining if text]
    if len(non_empty_remaining) > 1:
        return False
    return True


def looks_like_date(text):
    return bool(DATE_RE.match(str(text).strip()))


def looks_like_secondary_item(text):
    value = str(text).strip()
    if not value:
        return False
    if looks_like_date(value):
        return True
    if value.endswith((":", "：")):
        return False
    if any(token in value for token in ("基线", "术后", "(", "（", "kg", "岁", "月")):
        return False
    if len(value) <= 12:
        return True
    return False


def classify_rows(table):
    roles = ["header"]
    in_block = False
    in_inline = False

    for row_index in range(1, len(table.rows)):
        row = table.rows[row_index]
        first_text = row.cells[0].text.strip()

        if is_group_heading_row(row):
            roles.append("block_heading")
            in_block = True
            in_inline = False
            continue

        if is_inline_heading_row(row):
            roles.append("inline_heading")
            in_inline = True
            continue

        if in_inline and looks_like_secondary_item(first_text):
            roles.append("child_secondary")
            continue

        in_inline = False
        if in_block:
            roles.append("child_primary")
        else:
            roles.append("body")

    return roles


def infer_alignment(text, header=False, group_heading=False):
    if header:
        return WD_ALIGN_PARAGRAPH.CENTER
    if group_heading:
        return WD_ALIGN_PARAGRAPH.LEFT
    value = str(text).strip()
    if not value:
        return WD_ALIGN_PARAGRAPH.LEFT
    if NUMERIC_RE.match(value) or value in ("Ref.", "Ref"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def set_row_property(row, tag_name):
    tr_pr = row._tr.get_or_add_trPr()
    element = tr_pr.find(qn("w:%s" % tag_name))
    if element is None:
        element = OxmlElement("w:%s" % tag_name)
        tr_pr.append(element)
    return element


def set_repeat_header(row):
    element = set_row_property(row, "tblHeader")
    element.set(qn("w:val"), "true")


def set_cant_split(row):
    set_row_property(row, "cantSplit")


def estimate_table_width_cm(table):
    if not table.rows:
        return 0.0
    column_count = len(table.rows[0].cells)
    if column_count == 0:
        return 0.0

    estimated = 0.0
    sample_rows = table.rows[: min(len(table.rows), 40)]
    for column_index in range(column_count):
        max_len = 0
        for row in sample_rows:
            try:
                text = row.cells[column_index].text.strip()
            except IndexError:
                text = ""
            max_len = max(max_len, len(text))
        estimated += max(1.8, min(5.2, 0.16 * max_len + 0.8))
    return estimated


def table_needs_landscape(table):
    if not table.rows:
        return False
    column_count = len(table.rows[0].cells)
    if column_count >= 6:
        return True
    return estimate_table_width_cm(table) > PORTRAIT_CONTENT_WIDTH_CM


def apply_three_line_style(table, western_font, east_asia_body, east_asia_heading, font_size):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    clear_borders(table)
    row_roles = classify_rows(table)

    for row_index, row in enumerate(table.rows):
        header = row_index == 0
        role = row_roles[row_index]
        block_heading = role == "block_heading"
        inline_heading = role == "inline_heading"
        set_cant_split(row)
        if header:
            set_repeat_header(row)
        for cell_index, cell in enumerate(unique_cells(row.cells)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            east_font = east_asia_body
            if header or block_heading:
                east_font = east_asia_heading
            elif inline_heading and cell_index == 0:
                east_font = east_asia_heading
            bold = header or block_heading or (inline_heading and cell_index == 0)
            left_indent = 0.0
            if cell_index == 0:
                if role == "child_primary":
                    left_indent = 0.5
                elif role == "inline_heading":
                    left_indent = 0.5
                elif role == "child_secondary":
                    left_indent = 1.0
            for paragraph in cell.paragraphs:
                style_paragraph(
                    paragraph,
                    western_font,
                    east_font,
                    font_size,
                    bold=bold,
                    italic=False,
                    alignment=infer_alignment(
                        paragraph.text,
                        header=header,
                        group_heading=(block_heading or inline_heading),
                    ),
                    left_indent_cm=left_indent,
                )

    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for cell in unique_cells(table.rows[0].cells):
        set_cell_border(cell, top=line, bottom=line)
    for cell in unique_cells(table.rows[-1].cells):
        set_cell_border(cell, bottom=line)


def apply_page_layout(document, force_landscape=False):
    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        if force_landscape:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)


def get_body_sectpr(document):
    body = document._body._element
    for child in reversed(list(body.iterchildren())):
        if child.tag.endswith("}sectPr"):
            return child
    raise ValueError("Document body does not contain sectPr")


def set_or_update_child(parent, tag_name):
    child = parent.find(qn("w:%s" % tag_name))
    if child is None:
        child = OxmlElement("w:%s" % tag_name)
        parent.append(child)
    return child


def configure_sectpr(sectpr, landscape=False, continuous=True):
    type_el = set_or_update_child(sectpr, "type")
    if continuous:
        type_el.set(qn("w:val"), "continuous")
    else:
        if qn("w:val") in type_el.attrib:
            del type_el.attrib[qn("w:val")]

    pg_sz = set_or_update_child(sectpr, "pgSz")
    width, height = A4_LANDSCAPE if landscape else A4_PORTRAIT
    pg_sz.set(qn("w:w"), str(width))
    pg_sz.set(qn("w:h"), str(height))
    if landscape:
        pg_sz.set(qn("w:orient"), "landscape")
    elif qn("w:orient") in pg_sz.attrib:
        del pg_sz.attrib[qn("w:orient")]

    pg_mar = set_or_update_child(sectpr, "pgMar")
    for edge in ("top", "right", "bottom", "left"):
        pg_mar.set(qn("w:%s" % edge), str(MARGIN_TWIPS))
    pg_mar.set(qn("w:header"), "720")
    pg_mar.set(qn("w:footer"), "720")
    pg_mar.set(qn("w:gutter"), "0")


def apply_wide_table_sections(document, table_infos, args):
    body_sectpr = get_body_sectpr(document)
    configure_sectpr(body_sectpr, landscape=False, continuous=True)
    use_continuous = args.wide_table_section_mode == "continuous"

    for info in table_infos:
        if not info["needs_landscape"]:
            continue

        start_block = info["caption"] or info["table"]
        end_block = info["note"] or info["table"]

        start_break = insert_paragraph_before_block(start_block)
        style_paragraph(
            start_break,
            args.font_name,
            args.east_asia_font,
            args.note_size,
            bold=False,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=args.wide_table_spacing_before_pt,
            space_after=0,
            line_spacing=1.0,
        )
        start_ppr = start_break._p.get_or_add_pPr()
        start_sectpr = OxmlElement("w:sectPr")
        configure_sectpr(
            start_sectpr,
            landscape=False,
            continuous=use_continuous,
        )
        start_ppr.append(start_sectpr)

        end_break = insert_paragraph_after_block(end_block)
        style_paragraph(
            end_break,
            args.font_name,
            args.east_asia_font,
            args.note_size,
            bold=False,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=0,
            space_after=args.wide_table_spacing_after_pt,
            line_spacing=1.0,
        )
        end_ppr = end_break._p.get_or_add_pPr()
        end_sectpr = OxmlElement("w:sectPr")
        configure_sectpr(
            end_sectpr,
            landscape=True,
            continuous=use_continuous,
        )
        end_ppr.append(end_sectpr)


def body_blocks(document):
    paragraphs_by_id = {id(p._p): p for p in document.paragraphs}
    tables_by_id = {id(t._tbl): t for t in document.tables}
    blocks = []
    for child in document._body._element.iterchildren():
        if child.tag.endswith("}p") and id(child) in paragraphs_by_id:
            blocks.append(("paragraph", paragraphs_by_id[id(child)]))
        elif child.tag.endswith("}tbl") and id(child) in tables_by_id:
            blocks.append(("table", tables_by_id[id(child)]))
    return blocks


def style_related_paragraphs(blocks, western_font, east_asia_body, east_asia_heading, caption_size, note_size):
    table_number = 0
    table_infos = []
    for index, (kind, obj) in enumerate(blocks):
        if kind != "table":
            continue
        table_number += 1
        caption_paragraph = None
        if index > 0 and blocks[index - 1][0] == "paragraph":
            prev_paragraph = blocks[index - 1][1]
            if TABLE_CAPTION_RE.match(prev_paragraph.text.strip()):
                prev_paragraph.text = normalize_caption_text(
                    "Table",
                    table_number,
                    prev_paragraph.text,
                    TABLE_CAPTION_RE,
                )
                caption_paragraph = prev_paragraph
        if caption_paragraph is None:
            caption_paragraph = insert_paragraph_before_table(obj, "Table %s." % table_number)
        style_paragraph(
            caption_paragraph,
            western_font,
            east_asia_heading,
            caption_size,
            bold=True,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6,
            space_after=6,
        )
        note_paragraph = None
        if index + 1 < len(blocks) and blocks[index + 1][0] == "paragraph":
            next_paragraph = blocks[index + 1][1]
            if NOTE_RE.match(next_paragraph.text.strip()):
                note_paragraph = next_paragraph
                style_paragraph(
                    next_paragraph,
                    western_font,
                    east_asia_body,
                    note_size,
                    bold=False,
                    italic=False,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=6,
                    space_after=6,
                )
        table_infos.append(
            {
                "table": obj,
                "caption": caption_paragraph,
                "note": note_paragraph,
                "needs_landscape": table_needs_landscape(obj),
            }
        )
    return table_infos


def main():
    args = parse_args()
    document = Document(args.source)
    force_landscape = args.landscape
    apply_page_layout(document, force_landscape=force_landscape)
    for table in document.tables:
        apply_three_line_style(
            table,
            args.font_name,
            args.east_asia_font,
            args.east_asia_heading_font,
            args.font_size,
        )
    table_infos = style_related_paragraphs(
        body_blocks(document),
        args.font_name,
        args.east_asia_font,
        args.east_asia_heading_font,
        args.caption_size,
        args.note_size,
    )
    if not force_landscape:
        apply_wide_table_sections(document, table_infos, args)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print("Wrote %s" % output_path)


if __name__ == "__main__":
    main()
