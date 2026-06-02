#!/usr/bin/env python3
"""Convert CSV/XLSX data into a grouped three-line table in DOCX."""

import argparse
import json
import re
import sys
from collections import OrderedDict
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

LIKELY_GROUP_COLUMNS = (
    "group",
    "category",
    "section",
    "domain",
    "block",
    "heading",
)
DEFAULT_GROUPING_RULES_PATH = (
    Path(__file__).resolve().parent.parent / "assets" / "medical-grouping-rules.json"
)


def parse_args():
    parser = argparse.ArgumentParser(
        description=(
            "Convert CSV/XLSX into a DOCX three-line table with optional "
            "journal-style grouped section headings."
        )
    )
    parser.add_argument("source", help="Path to .csv or .xlsx input data")
    parser.add_argument("output", help="Path to output .docx")
    parser.add_argument(
        "--input-docx",
        help="Optional existing .docx to append the generated table to",
    )
    parser.add_argument("--sheet", help="Worksheet name for .xlsx input")
    parser.add_argument("--caption", help="Caption to place above the table")
    parser.add_argument("--note", help="Note to place below the table")
    parser.add_argument(
        "--label-column",
        help="Column containing row labels or variable names; defaults to the first non-group column",
    )
    parser.add_argument(
        "--group-column",
        help="Column whose values are section headings such as Demographics or Outcomes",
    )
    parser.add_argument(
        "--group-map",
        help=(
            "JSON file mapping section headings to row labels. "
            'Example: {"Demographics": ["Sex", "Age"], "Outcomes": ["Mortality"]}'
        ),
    )
    parser.add_argument(
        "--group-rules",
        default=str(DEFAULT_GROUPING_RULES_PATH),
        help=(
            "JSON file defining fixed medical grouping rules for automatic heading inference; "
            "defaults to assets/medical-grouping-rules.json"
        ),
    )
    parser.add_argument(
        "--drop-group-column",
        action="store_true",
        help="Remove the group column from the rendered table",
    )
    parser.add_argument(
        "--font-name",
        default="Times New Roman",
        help="Font name for table text; defaults to Times New Roman",
    )
    parser.add_argument(
        "--east-asia-font",
        default="宋体",
        help="East Asian body font; defaults to 宋体",
    )
    parser.add_argument(
        "--east-asia-heading-font",
        default="黑体",
        help="East Asian heading font for captions and grouped headings; defaults to 黑体",
    )
    parser.add_argument(
        "--font-size",
        default=11.0,
        type=float,
        help="Font size in points; defaults to 11.0",
    )
    parser.add_argument(
        "--caption-size",
        default=11.0,
        type=float,
        help="Caption font size in points; defaults to 11.0",
    )
    parser.add_argument(
        "--note-size",
        default=10.0,
        type=float,
        help="Note font size in points; defaults to 10.0",
    )
    parser.add_argument(
        "--landscape",
        action="store_true",
        help="Render the section in landscape orientation",
    )
    return parser.parse_args()


def error(message):
    print("ERROR: %s" % message, file=sys.stderr)
    sys.exit(1)


def normalize(text):
    chars = []
    for ch in str(text).strip().lower():
        if ch.isalnum():
            chars.append(ch)
    return "".join(chars)


def text_lower(value):
    return str(value).strip().lower()


def contains_any(text, keywords):
    lowered = text_lower(text)
    return any(keyword in lowered for keyword in keywords)


def looks_like_date(label):
    return bool(re.match(r"^\d{4}-\d{2}-\d{2}$", str(label).strip()))


def load_grouping_rules(path):
    with open(path, "r", encoding="utf-8") as handle:
        raw = json.load(handle)
    groups = raw.get("groups", [])
    for group in groups:
        group["keywords"] = [str(item).lower() for item in group.get("keywords", [])]
    raw["groups"] = groups
    raw["followup_markers"] = [str(item) for item in raw.get("followup_markers", [])]
    raw["outcome_keywords"] = [
        str(item).lower() for item in raw.get("outcome_keywords", [])
    ]
    raw["outcome_forced_labels"] = [
        str(item) for item in raw.get("outcome_forced_labels", [])
    ]
    raw["preserve_previous_suffixes"] = [
        str(item) for item in raw.get("preserve_previous_suffixes", [":", "："])
    ]
    raw["default_group"] = str(raw.get("default_group", "其他项目"))
    raw["followup_group"] = str(raw.get("followup_group", "随访指标"))
    raw["outcome_group"] = str(raw.get("outcome_group", "结局事件"))
    return raw


def infer_group_title(label, grouping_rules, previous_group=None):
    text = str(label).strip()
    if not text:
        return previous_group or grouping_rules["default_group"]

    followup = None
    for marker in grouping_rules["followup_markers"]:
        if marker in text:
            followup = marker
            break
    if looks_like_date(text) or any(
        forced_label in text for forced_label in grouping_rules["outcome_forced_labels"]
    ):
        return grouping_rules["outcome_group"]
    if followup and contains_any(text, grouping_rules["outcome_keywords"]):
        return grouping_rules["outcome_group"]
    if followup:
        return grouping_rules["followup_group"]
    for group in grouping_rules["groups"]:
        if contains_any(text, group["keywords"]):
            return group["title"]
    if contains_any(text, grouping_rules["outcome_keywords"]):
        return grouping_rules["outcome_group"]
    if text.endswith(tuple(grouping_rules["preserve_previous_suffixes"])) and previous_group:
        return previous_group
    return previous_group or grouping_rules["default_group"]


def load_frame(path, sheet_name=None):
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".csv":
        frame = pd.read_csv(source)
    elif suffix in (".xlsx", ".xlsm", ".xls"):
        frame = pd.read_excel(source, sheet_name=sheet_name)
    else:
        error("Unsupported source file type: %s" % source.suffix)

    if isinstance(frame, dict):
        if not frame:
            error("No worksheets were found in the Excel file.")
        first_name = next(iter(frame))
        frame = frame[first_name]

    frame = frame.fillna("")
    frame.columns = [str(col).strip() for col in frame.columns]
    if len(frame.columns) > 0:
        first_column = frame.columns[0]
        if str(first_column).lower().startswith("unnamed:"):
            frame = frame.rename(columns={first_column: "项目"})
    return frame


def choose_group_column(frame, requested):
    if requested:
        if requested not in frame.columns:
            error("Requested group column not found: %s" % requested)
        return requested

    normalized_columns = OrderedDict(
        (normalize(column), column) for column in frame.columns
    )
    for candidate in LIKELY_GROUP_COLUMNS:
        matched = normalized_columns.get(normalize(candidate))
        if matched:
            return matched
    return None


def choose_label_column(frame, requested, group_column):
    if requested:
        if requested not in frame.columns:
            error("Requested label column not found: %s" % requested)
        return requested

    for column in frame.columns:
        if column != group_column:
            return column
    error("Could not determine a label column.")


def load_group_map(path):
    with open(path, "r", encoding="utf-8") as handle:
        data = json.load(handle, object_pairs_hook=OrderedDict)
    if not isinstance(data, dict) or not data:
        error("Group map must be a non-empty JSON object.")

    prepared = OrderedDict()
    for group_name, labels in data.items():
        if isinstance(labels, str):
            labels = [labels]
        if not isinstance(labels, list) or not labels:
            error("Each group map entry must be a non-empty list or string.")
        prepared[str(group_name)] = set(normalize(label) for label in labels)
    return prepared


def rows_to_blocks(rows, group_names):
    blocks = []
    current_name = None
    current_rows = []
    for row, group_name in zip(rows, group_names):
        if group_name != current_name:
            if current_rows:
                blocks.append((current_name, current_rows))
            current_name = group_name
            current_rows = [row]
        else:
            current_rows.append(row)
    if current_rows:
        blocks.append((current_name, current_rows))
    return blocks


def assign_groups(frame, label_column, group_column, group_map, grouping_rules):
    rows = frame.to_dict(orient="records")
    if group_column:
        group_names = []
        for row in rows:
            group_name = str(row.get(group_column, "")).strip() or "Ungrouped"
            group_names.append(group_name)
        return rows_to_blocks(rows, group_names)

    if group_map:
        grouped_names = []
        for row in rows:
            label = normalize(row.get(label_column, ""))
            matched = False
            for group_name, group_labels in group_map.items():
                if label in group_labels:
                    grouped_names.append(group_name)
                    matched = True
                    break
            if not matched:
                grouped_names.append("Ungrouped")
        return rows_to_blocks(rows, grouped_names)

    inferred_names = []
    current_group = None
    for row in rows:
        current_group = infer_group_title(
            row.get(label_column, ""),
            grouping_rules,
            current_group,
        )
        inferred_names.append(current_group)
    if inferred_names and set(inferred_names) == {grouping_rules["default_group"]}:
        return [("", rows)]
    return rows_to_blocks(rows, inferred_names)


def create_document(input_docx=None, landscape=False):
    document = Document(input_docx) if input_docx else Document()
    section = document.sections[-1]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    return document


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:%s" % key), str(value))


def clear_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                top={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
            )


def set_run_fonts(run, western_font, east_asia_font, size, bold=False, italic=False):
    run.font.name = western_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western_font)
    r_fonts.set(qn("w:hAnsi"), western_font)
    r_fonts.set(qn("w:cs"), western_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def style_cell_paragraph(
    paragraph,
    western_font,
    east_asia_font,
    font_size,
    bold=False,
    italic=False,
    alignment=None,
):
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_fonts(
            run,
            western_font,
            east_asia_font,
            font_size,
            bold=bold,
            italic=italic,
        )


def style_free_paragraph(
    paragraph,
    western_font,
    east_asia_font,
    font_size,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent_cm=0.0,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_fonts(
            run,
            western_font,
            east_asia_font,
            font_size,
            bold=bold,
            italic=italic,
        )


def infer_alignment(value):
    text = str(value).strip()
    if text == "":
        return WD_ALIGN_PARAGRAPH.LEFT
    try:
        float(text.replace(",", ""))
        return WD_ALIGN_PARAGRAPH.RIGHT
    except ValueError:
        return WD_ALIGN_PARAGRAPH.LEFT


def add_group_row(
    table,
    text,
    column_count,
    western_font,
    east_asia_body_font,
    east_asia_heading_font,
    font_size,
):
    cells = table.add_row().cells
    merged = cells[0]
    for index in range(1, column_count):
        merged = merged.merge(cells[index])
    merged.text = text
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = merged.paragraphs[0]
    style_cell_paragraph(
        paragraph,
        western_font,
        east_asia_heading_font or east_asia_body_font,
        font_size,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )


def add_data_row(table, values, headers, western_font, east_asia_body_font, font_size):
    row = table.add_row()
    for index, header in enumerate(headers):
        cell = row.cells[index]
        cell.text = str(values.get(header, ""))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        style_cell_paragraph(
            paragraph,
            western_font,
            east_asia_body_font,
            font_size,
            bold=False,
            alignment=infer_alignment(values.get(header, "")),
        )


def apply_three_line_borders(table):
    clear_borders(table)
    header_row = table.rows[0]
    last_row = table.rows[-1]
    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}

    for cell in header_row.cells:
        set_cell_border(cell, top=line, bottom=line)
    for cell in last_row.cells:
        set_cell_border(cell, bottom=line)


def build_table(
    document,
    grouped_rows,
    headers,
    caption,
    note,
    western_font,
    east_asia_body_font,
    east_asia_heading_font,
    font_size,
    caption_size,
    note_size,
):
    if caption:
        paragraph = document.add_paragraph()
        style_free_paragraph(
            paragraph,
            western_font,
            east_asia_heading_font or east_asia_body_font,
            caption_size,
            bold=True,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_cm=0.0,
        )
        paragraph.runs[0].text = caption

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        style_cell_paragraph(
            header_cells[index].paragraphs[0],
            western_font,
            east_asia_heading_font or east_asia_body_font,
            font_size,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for group_name, rows in grouped_rows:
        if group_name:
            add_group_row(
                table,
                group_name,
                len(headers),
                western_font,
                east_asia_body_font,
                east_asia_heading_font,
                font_size,
            )
        for row in rows:
            add_data_row(
                table,
                row,
                headers,
                western_font,
                east_asia_body_font,
                font_size,
            )

    apply_three_line_borders(table)

    if note:
        note_paragraph = document.add_paragraph()
        style_free_paragraph(
            note_paragraph,
            western_font,
            east_asia_body_font,
            note_size,
            bold=False,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_cm=0.0,
        )
        note_paragraph.runs[0].text = "Note. %s" % note


def main():
    args = parse_args()
    frame = load_frame(args.source, args.sheet)
    group_column = choose_group_column(frame, args.group_column)
    label_column = choose_label_column(frame, args.label_column, group_column)
    group_map = load_group_map(args.group_map) if args.group_map else None
    grouping_rules = load_grouping_rules(args.group_rules)
    grouped_rows = assign_groups(
        frame,
        label_column,
        group_column,
        group_map,
        grouping_rules,
    )

    headers = list(frame.columns)
    if args.drop_group_column and group_column:
        headers = [column for column in headers if column != group_column]
        for _, rows in grouped_rows:
            for row in rows:
                row.pop(group_column, None)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = create_document(args.input_docx, args.landscape)
    build_table(
        document,
        grouped_rows,
        headers,
        args.caption,
        args.note,
        args.font_name,
        args.east_asia_font,
        args.east_asia_heading_font,
        args.font_size,
        args.caption_size,
        args.note_size,
    )
    document.save(str(output_path))
    print("Wrote %s" % output_path)


if __name__ == "__main__":
    main()
