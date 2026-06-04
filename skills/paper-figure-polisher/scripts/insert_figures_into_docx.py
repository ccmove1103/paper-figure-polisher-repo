#!/usr/bin/env python3
"""Insert new image figures into a DOCX while preserving table-before-figure order."""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


FIGURE_CAPTION_RE = re.compile(r"^(figure|图)\s*(\d+)\s*[\.:\-：]?\s*", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(r"^(table|表)\s*(\d+)\s*[\.:\-：]?\s*", re.IGNORECASE)
NOTE_RE = re.compile(r"^(note\.|注[:：])", re.IGNORECASE)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Insert one or more image figures into an existing DOCX."
    )
    parser.add_argument("source_docx", help="Path to source .docx")
    parser.add_argument("output", help="Path to output .docx")
    parser.add_argument("images", nargs="+", help="Image files to insert")
    parser.add_argument(
        "--captions",
        nargs="*",
        default=[],
        help="Optional figure caption tails aligned with the image list",
    )
    parser.add_argument(
        "--notes",
        nargs="*",
        default=[],
        help="Optional figure notes aligned with the image list",
    )
    parser.add_argument(
        "--width-cm",
        default=12.0,
        type=float,
        help="Inserted figure width in centimeters; defaults to 12",
    )
    parser.add_argument(
        "--font-name",
        default="Times New Roman",
        help="Western font for captions; defaults to Times New Roman",
    )
    parser.add_argument(
        "--east-asia-font",
        default="宋体",
        help="East Asian body font; defaults to 宋体",
    )
    parser.add_argument(
        "--east-asia-heading-font",
        default="黑体",
        help="East Asian heading font; defaults to 黑体",
    )
    parser.add_argument(
        "--caption-size",
        default=11.0,
        type=float,
        help="Figure caption size in points; defaults to 11",
    )
    parser.add_argument(
        "--note-size",
        default=10.0,
        type=float,
        help="Figure note size in points; defaults to 10",
    )
    return parser.parse_args()


def has_drawing(paragraph):
    for node in paragraph._p.iter():
        if node.tag.endswith("}drawing") or node.tag.endswith("}pict"):
            return True
    return False


def body_blocks(document):
    paragraphs_by_id = {id(p._p): p for p in document.paragraphs}
    tables_by_id = {id(t._tbl): t for t in document.tables}
    blocks = []
    for child in document._body._element.iterchildren():
        if child.tag.endswith("}p") and id(child) in paragraphs_by_id:
            blocks.append(("paragraph", paragraphs_by_id[id(child)]))
        elif child.tag.endswith("}tbl") and id(child) in tables_by_id:
            blocks.append(("table", tables_by_id[id(child)]))
    return blocks


def block_element(block):
    return block._p if hasattr(block, "_p") else block._tbl


def next_index_if(blocks, index, predicate):
    candidate = index + 1
    if candidate < len(blocks) and predicate(blocks[candidate]):
        return candidate
    return index


def table_block_end(blocks, table_index):
    end_index = table_index
    while end_index + 1 < len(blocks):
        kind, obj = blocks[end_index + 1]
        if kind != "paragraph" or not NOTE_RE.match(obj.text.strip()):
            break
        end_index += 1
    return end_index


def figure_block_end(blocks, figure_index):
    end_index = figure_index
    end_index = next_index_if(
        blocks,
        end_index,
        lambda item: item[0] == "paragraph"
        and FIGURE_CAPTION_RE.match(item[1].text.strip()),
    )
    end_index = next_index_if(
        blocks,
        end_index,
        lambda item: item[0] == "paragraph" and NOTE_RE.match(item[1].text.strip()),
    )
    if end_index == figure_index:
        end_index = next_index_if(
            blocks,
            end_index,
            lambda item: item[0] == "paragraph" and NOTE_RE.match(item[1].text.strip()),
        )
    return end_index


def choose_insertion_target(document):
    blocks = body_blocks(document)
    if not blocks:
        return "append", None

    last_figure_end = None
    last_table_end = None
    for index, (kind, obj) in enumerate(blocks):
        if kind == "paragraph" and has_drawing(obj):
            last_figure_end = figure_block_end(blocks, index)
        elif kind == "table":
            last_table_end = table_block_end(blocks, index)

    if last_figure_end is not None:
        return "after", blocks[last_figure_end][1]
    if last_table_end is not None:
        return "after", blocks[last_table_end][1]
    return "append", None


def relocate_blocks(blocks, placement, anchor):
    if placement == "append" or anchor is None:
        return

    current_element = block_element(anchor)
    for block in blocks:
        moved = block_element(block)
        current_element.addnext(moved)
        current_element = moved


def set_run_fonts(run, western_font, east_asia_font, size, bold=False, italic=False):
    run.font.name = western_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western_font)
    r_fonts.set(qn("w:hAnsi"), western_font)
    r_fonts.set(qn("w:cs"), western_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def style_paragraph(
    paragraph,
    western_font,
    east_asia_font,
    size,
    bold=False,
    italic=False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_fonts(
            run,
            western_font,
            east_asia_font,
            size,
            bold=bold,
            italic=italic,
        )


def count_existing_figures(document):
    return sum(1 for paragraph in document.paragraphs if has_drawing(paragraph))


def build_caption_text(number, caption_tail):
    tail = str(caption_tail or "").strip()
    if not tail:
        return "Figure %s." % number
    if FIGURE_CAPTION_RE.match(tail):
        tail = FIGURE_CAPTION_RE.sub("", tail).strip()
    return "Figure %s.%s" % (number, tail)


def build_note_text(note):
    text = str(note or "").strip()
    if not text:
        return ""
    if NOTE_RE.match(text):
        return text
    return "Note. %s" % text


def add_figure_blocks(document, image_path, caption_text, note_text, args):
    created = []

    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.space_before = Pt(6)
    figure_paragraph.paragraph_format.space_after = Pt(6)
    figure_paragraph.paragraph_format.line_spacing = 1.5
    figure_paragraph.add_run().add_picture(str(image_path), width=Cm(args.width_cm))
    created.append(figure_paragraph)

    caption_paragraph = document.add_paragraph(caption_text)
    style_paragraph(
        caption_paragraph,
        args.font_name,
        args.east_asia_heading_font,
        args.caption_size,
        bold=False,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    created.append(caption_paragraph)

    if note_text:
        note_paragraph = document.add_paragraph(note_text)
        style_paragraph(
            note_paragraph,
            args.font_name,
            args.east_asia_font,
            args.note_size,
            bold=False,
            italic=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        created.append(note_paragraph)

    return created


def main():
    args = parse_args()
    document = Document(args.source_docx)
    placement, anchor = choose_insertion_target(document)
    existing_count = count_existing_figures(document)

    created_blocks = []
    for offset, image in enumerate(args.images, start=1):
        image_path = Path(image)
        caption_tail = args.captions[offset - 1] if offset - 1 < len(args.captions) else ""
        note_value = args.notes[offset - 1] if offset - 1 < len(args.notes) else ""
        caption_text = build_caption_text(existing_count + offset, caption_tail)
        note_text = build_note_text(note_value)
        created_blocks.extend(
            add_figure_blocks(document, image_path, caption_text, note_text, args)
        )

    relocate_blocks(created_blocks, placement, anchor)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print("Wrote %s" % output_path)


if __name__ == "__main__":
    main()
